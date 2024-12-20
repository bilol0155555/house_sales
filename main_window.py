import xlsxwriter
from PyQt6.QtGui import QPixmap, QIcon
from PyQt6.QtWidgets import QMainWindow, QVBoxLayout, QPushButton, QWidget, QMessageBox, QTableWidget, QTableWidgetItem, \
    QInputDialog, QLabel, QFileDialog, QHBoxLayout, QLineEdit, QDialog, QCheckBox
from PyQt6.QtCore import Qt, QSize, QFile, QTextStream
from docx import Document


from view_image import ImageViewer

class MainWindow(QMainWindow):
    def __init__(self, db, user):
        super().__init__()
        self.db = db
        self.user_id, self.role = user
        self.setWindowTitle("Продажа жилья")
        self.is_dark_theme = False

        layout = QVBoxLayout()

        self.theme_button = QPushButton("", self)
        self.theme_button.setFixedSize(40, 40)
        self.theme_button.clicked.connect(self.toggle_theme)
        layout.addWidget(self.theme_button, alignment=Qt.AlignmentFlag.AlignRight)

        # Поле для поиска
        search_layout = QHBoxLayout()
        self.search_input = QLineEdit(self)
        self.search_input.setPlaceholderText("Поиск по адресу или цене")
        search_layout.addWidget(self.search_input)

        search_button = QPushButton("Поиск", self)
        search_button.setIcon(QIcon("icons/search.png"))
        search_button.clicked.connect(self.search_properties)
        search_layout.addWidget(search_button)

        layout.addLayout(search_layout)

        if self.role == "admin":
            self.init_admin_ui(layout)
        else:
            self.init_user_ui(layout)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)
        self.apply_stylesheet()

    def apply_stylesheet(self):
        file = QFile("styles/light_theme.css")
        if not file.exists():
            QMessageBox.warning(self, "Ошибка", "Файл стилей не найден")
            return
        file.open(QFile.OpenModeFlag.ReadOnly)
        stylesheet = QTextStream(file).readAll()
        self.setStyleSheet(stylesheet)

    def toggle_theme(self):
        self.is_dark_theme = not self.is_dark_theme
        theme_file = "styles/dark_theme.css" if self.is_dark_theme else "styles/light_theme.css"

        file = QFile(theme_file)
        if not file.exists():
            QMessageBox.warning(self, "Ошибка", "Файл стилей не найден")
            return
        file.open(QFile.OpenModeFlag.ReadOnly)
        stylesheet = QTextStream(file).readAll()
        self.setStyleSheet(stylesheet)

        icon = QIcon("icons/moon.png") if self.is_dark_theme else QIcon("icons/sun.png")
        self.theme_button.setIcon(icon)

    def init_admin_ui(self, layout):
        self.properties_table = QTableWidget(self)
        self.refresh_properties()
        layout.addWidget(self.properties_table)

        add_button = QPushButton("Добавить жилье", self)
        add_button.setIcon(QIcon("icons/add.png"))
        add_button.clicked.connect(self.add_property)
        layout.addWidget(add_button)

        edit_button = QPushButton("Редактировать жилье", self)
        edit_button.setIcon(QIcon("icons/edit.png"))
        edit_button.clicked.connect(self.edit_property)
        layout.addWidget(edit_button)

        delete_button = QPushButton("Удалить жилье", self)
        delete_button.setIcon(QIcon("icons/delete.png"))
        delete_button.clicked.connect(self.delete_property)
        layout.addWidget(delete_button)

        # Переключатель уведомлений
        notification_checkbox = QCheckBox("Уведомлять о изменениях")
        notification_checkbox.setChecked(True)
        notification_checkbox.stateChanged.connect(self.toggle_notifications)
        layout.addWidget(notification_checkbox)




        export_excel_button = QPushButton("Экспорт в Excel", self)
        export_excel_button.setIcon(QIcon("icons/excel.png"))
        export_excel_button.clicked.connect(self.export_to_excel)
        layout.addWidget(export_excel_button)

        export_word_button = QPushButton("Экспорт в Word", self)
        export_word_button.setIcon(QIcon("icons/word.png"))
        export_word_button.clicked.connect(self.export_to_word)
        layout.addWidget(export_word_button)

    def init_user_ui(self, layout):
        self.properties_table = QTableWidget(self)
        self.refresh_properties()
        layout.addWidget(self.properties_table)

        view_button = QPushButton("Просмотреть", self)
        view_button.setIcon(QIcon("icons/view.png"))
        view_button.clicked.connect(self.view_property)
        layout.addWidget(view_button)

        purchase_button = QPushButton("Купить", self)
        purchase_button.clicked.connect(self.purchase_property)
        layout.addWidget(purchase_button)

        history_button = QPushButton("История покупок", self)
        history_button.clicked.connect(self.view_purchase_history)
        layout.addWidget(history_button)

    def refresh_properties(self):
        cursor = self.db.conn.cursor()
        cursor.execute("SELECT id, address, price, rooms, area, owner FROM properties")
        properties = cursor.fetchall()

        self.properties_table.setRowCount(len(properties))
        self.properties_table.setColumnCount(7)
        self.properties_table.setHorizontalHeaderLabels(["ID", "Адрес", "Цена", "Комнат", "Площадь (м²)", "Владелец", "Изображения"])

        for row, (prop_id, address, price, rooms, area, owner) in enumerate(properties):
            self.properties_table.setItem(row, 0, QTableWidgetItem(str(prop_id)))
            self.properties_table.setItem(row, 1, QTableWidgetItem(address))
            self.properties_table.setItem(row, 2, QTableWidgetItem(f"{price:.2f}"))
            self.properties_table.setItem(row, 3, QTableWidgetItem(str(rooms)))
            self.properties_table.setItem(row, 4, QTableWidgetItem(f"{area:.2f} м²"))
            self.properties_table.setItem(row, 5, QTableWidgetItem(owner or "Не указан"))



    def toggle_notifications(self, state):
        self.notifications_enabled = (state == Qt.CheckState.Checked)
        if self.notifications_enabled:
            self.show_notification("Уведомления включены")
        else:
            self.show_notification("Уведомления выключены")



    def show_notification(self, message):
        # Пример уведомления
        print(f"Notification: {message}")


    def add_property(self):
        address, ok1 = QInputDialog.getText(self, "Добавить жилье", "Адрес:")
        if not ok1 or not address:
            return

        price, ok2 = QInputDialog.getDouble(self, "Добавить жилье", "Цена:")
        if not ok2:
            return

        rooms, ok3 = QInputDialog.getInt(self, "Добавить жилье", "Количество комнат:")
        if not ok3:
            return

        area, ok4 = QInputDialog.getDouble(self, "Добавить жилье", "Площадь (м²):")
        if not ok4:
            return

        owner, ok5 = QInputDialog.getText(self, "Добавить жилье", "Владелец:")
        if not ok5:
            owner = None

        with self.db.conn:
            self.db.conn.execute(
                "INSERT INTO properties (address, price, rooms, area, owner) VALUES (?, ?, ?, ?, ?)",
                (address, price, rooms, area, owner),
            )
        QMessageBox.information(self, "Успех", "Жилье добавлено")
        self.refresh_properties()

    def export_to_excel(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить как",
                                                   "", "Excel Files (*.xlsx)")
        if not file_path:
            return

        workbook = xlsxwriter.Workbook(file_path)
        worksheet = workbook.add_worksheet()

        headers = ["ID", "Адрес", "Цена", "Комнат", "Площадь (м²)", "Владелец"]
        for col, header in enumerate(headers):
            worksheet.write(0, col, header)

        for row in range(self.properties_table.rowCount()):
            for col in range(self.properties_table.columnCount()):
                item = self.properties_table.item(row, col)
                worksheet.write(row + 1, col, item.text() if item else "")

        workbook.close()
        QMessageBox.information(self, "Экспорт завершен", "Данные успешно сохранены в Excel")

    def export_to_word(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить как", "", "Word Files (*.docx)")
        if not file_path:
            return

        document = Document()
        document.add_heading("Список недвижимости", level=1)

        table = document.add_table(rows=1, cols=6)
        headers = ["ID", "Адрес", "Цена", "Комнат", "Площадь (м²)", "Владелец"]
        hdr_cells = table.rows[0].cells
        for col, header in enumerate(headers):
            hdr_cells[col].text = header

        for row in range(self.properties_table.rowCount()):
            cells = table.add_row().cells
            for col in range(self.properties_table.columnCount()):
                item = self.properties_table.item(row, col)
                cells[col].text = item.text() if item else ""

        document.save(file_path)
        QMessageBox.information(self, "Экспорт завершен", "Данные успешно сохранены в Word")

    def view_property(self):
        selected_items = self.properties_table.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Ошибка", "Выберите жилье для просмотра")
            return

        prop_id = selected_items[0].text()
        cursor = self.db.conn.cursor()
        cursor.execute("SELECT images FROM properties WHERE id = ?", (prop_id,))
        result = cursor.fetchone()

        if not result or not result[0]:
            QMessageBox.warning(self, "Ошибка", "Нет доступных изображений")
            return

        images = result[0].split(";")
        try:
            image_viewer = ImageViewer(images)
            image_viewer.exec()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть изображения: {str(e)}")

    def edit_property(self):
        selected_row = self.properties_table.currentRow()
        if selected_row == -1:
            QMessageBox.warning(self, "Ошибка", "Выберите жилье для редактирования")
            return

        prop_id = int(self.properties_table.item(selected_row, 0).text())
        address = self.properties_table.item(selected_row, 1).text()
        price = float(self.properties_table.item(selected_row, 2).text())

        new_address, ok1 = QInputDialog.getText(self, "Редактировать жилье", "Новый адрес:",
                                                text=address)
        if not ok1 or not new_address:
            return

        new_price, ok2 = QInputDialog.getDouble(self, "Редактировать жилье", "Новая цена:",
                                                value=price)
        if not ok2:
            return

        image_path, _ = QFileDialog.getOpenFileName(self, "Выбрать изображение",
                                                    "", "Изображения (*.png *.jpg *.jpeg)")
        if not image_path:
            image_path = self.db.conn.execute("SELECT image FROM properties WHERE id = ?",
                                              (prop_id,)).fetchone()[0]

        with self.db.conn:
            self.db.conn.execute(
                "UPDATE properties SET address = ?, price = ?, image = ? WHERE id = ?",
                (new_address, new_price, image_path, prop_id),
            )
        QMessageBox.information(self, "Успех", "Жилье обновлено")
        self.refresh_properties()

    def delete_property(self):
        selected_row = self.properties_table.currentRow()
        if selected_row == -1:
            QMessageBox.warning(self, "Ошибка", "Выберите жилье для удаления")
            return

        prop_id = int(self.properties_table.item(selected_row, 0).text())
        reply = QMessageBox.question(
            self, "Удаление жилья", "Вы уверены, что хотите удалить выбранное жилье?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            with self.db.conn:
                self.db.conn.execute("DELETE FROM properties WHERE id = ?", (prop_id,))
            QMessageBox.information(self, "Успех", "Жилье удалено")
            self.refresh_properties()

    def view_property(self):
        selected_items = self.properties_table.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Ошибка", "Выберите жилье для просмотра")
            return

        prop_id = selected_items[0].text()
        cursor = self.db.conn.cursor()
        cursor.execute("SELECT image FROM properties WHERE id = ?", (prop_id,))
        result = cursor.fetchone()

        if not result or not result[0]:
            QMessageBox.warning(self, "Ошибка", "Нет доступных изображений")
            return

        images = result[0].split(";")
        try:
            image_viewer = ImageViewer(images)
            image_viewer.exec()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть изображения: {str(e)}")

    # def toggle_theme(self):
    #     self.is_dark_theme = not self.is_dark_theme
    #     theme_file = "styles/dark_theme.css" if self.is_dark_theme else "styles/light_theme.css"
    #
    #     file = QFile(theme_file)
    #     if not file.exists():
    #         QMessageBox.warning(self, "Ошибка", "Файл стилей не найден")
    #         return
    #     file.open(QFile.OpenModeFlag.ReadOnly)
    #     stylesheet = QTextStream(file).readAll()
    #     self.setStyleSheet(stylesheet)
    #
    #     icon = QIcon("icons/moon.png") if self.is_dark_theme else QIcon("icons/sun.png")
    #     self.theme_button.setIcon(icon)

    def search_properties(self):
        query = self.search_input.text()
        if not query.strip():
            self.refresh_properties()
            return

        cursor = self.db.conn.cursor()

        # Если запрос можно интерпретировать как число, ищем в столбце `price` как число
        try:
            numeric_query = float(query)
            cursor.execute(
                "SELECT id, address, price, image FROM properties WHERE address LIKE ? OR price = ?",
                (f"%{query}%", numeric_query)
            )
        except ValueError:
            # Если запрос не число, ищем только как строку
            cursor.execute(
                "SELECT id, address, price, image FROM properties WHERE address LIKE ? OR CAST(price AS TEXT) LIKE ?",
                (f"%{query}%", f"%{query}%")
            )

        properties = cursor.fetchall()

        self.properties_table.setRowCount(len(properties))
        for row, (prop_id, address, price, images) in enumerate(properties):
            self.properties_table.setItem(row, 0, QTableWidgetItem(str(prop_id)))
            self.properties_table.setItem(row, 1, QTableWidgetItem(address))

            # Форматирование цены с обработкой возможных ошибок
            try:
                formatted_price = f"{float(price):,.2f}"  # Преобразуем в число и форматируем
            except ValueError:
                formatted_price = str(price)  # Если не число, оставляем как есть

            self.properties_table.setItem(row, 2, QTableWidgetItem(formatted_price))

            if images:
                label = QLabel()
                pixmap = QPixmap(images.split(";")[0]).scaled(120, 80, Qt.AspectRatioMode.KeepAspectRatio)
                label.setPixmap(pixmap)
                label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                self.properties_table.setCellWidget(row, 7, label)
            else:
                self.properties_table.setItem(row, 7, QTableWidgetItem("Нет изображений"))

    def purchase_property(self):
        selected_items = self.properties_table.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Ошибка", "Выберите жилье для покупки")
            return

        prop_id = selected_items[0].text()
        cursor = self.db.conn.cursor()

        # Проверить, свободна ли квартира
        cursor.execute("SELECT id, address, price FROM properties WHERE id = ?", (prop_id,))
        property_info = cursor.fetchone()
        if not property_info:
            QMessageBox.warning(self, "Ошибка", "Выбранное жилье недоступно")
            return

        address, price = property_info[1], property_info[2]

        # Подтверждение покупки
        confirm = QMessageBox.question(
            self, "Подтверждение покупки",
            f"Вы уверены, что хотите купить жилье по адресу {address} за {price:.2f}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if confirm != QMessageBox.StandardButton.Yes:
            return

        # Запись покупки
        try:
            with self.db.conn:
                self.db.conn.execute(
                    "INSERT INTO purchase_history (user_id, property_id, address, price, purchase_date) VALUES (?, ?, ?, ?, datetime('now'))",
                    (self.user_id, prop_id, address, price)
                )
                self.db.conn.execute("DELETE FROM properties WHERE id = ?", (prop_id,))
            QMessageBox.information(self, "Успех", "Покупка успешно завершена!")
            self.refresh_properties()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось завершить покупку: {str(e)}")

    def view_purchase_history(self):
        cursor = self.db.conn.cursor()
        cursor.execute("""
            SELECT h.property_id, h.address, h.price, h.purchase_date
            FROM purchase_history h
            -- JOIN properties p ON h.property_id = p.id
            WHERE h.user_id = ?
        """, (self.user_id,))
        history = cursor.fetchall()
        print(history)

        if not history:
            QMessageBox.information(self, "История покупок", "У вас нет записей о покупках.")
            return

        history_window = QDialog(self)
        history_window.setWindowTitle("История покупок")
        history_window.setGeometry(300, 300, 600, 400)

        layout = QVBoxLayout()

        table = QTableWidget(len(history), 4)
        table.setHorizontalHeaderLabels(["ID", "Адрес", "Цена", "Дата покупки"])
        for row, (prop_id, address, price, purchase_date) in enumerate(history):
            table.setItem(row, 0, QTableWidgetItem(str(prop_id)))
            table.setItem(row, 1, QTableWidgetItem(address))
            table.setItem(row, 2, QTableWidgetItem(f"{price:.2f}"))
            table.setItem(row, 3, QTableWidgetItem(purchase_date))
        layout.addWidget(table)

        close_button = QPushButton("Закрыть", history_window)
        close_button.clicked.connect(history_window.close)
        layout.addWidget(close_button)

        history_window.setLayout(layout)
        history_window.exec()


